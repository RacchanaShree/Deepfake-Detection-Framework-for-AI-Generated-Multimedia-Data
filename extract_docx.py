from docx import Document

doc = Document('Proposal-2025-26.docx')
for para in doc.paragraphs:
    print(para.text)
    
# Also extract tables if any
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)
